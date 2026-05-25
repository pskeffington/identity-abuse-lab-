from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUTPUT_DIR = Path("paper/polished")
DOCX_PATH = OUTPUT_DIR / "Identity_Abuse_Lab_Polished_Report.docx"
PDF_PATH = OUTPUT_DIR / "Identity_Abuse_Lab_Polished_Report.pdf"
ZIP_PATH = OUTPUT_DIR / "Identity_Abuse_Lab_Polished_Deliverable.zip"
SCORED_EVENTS_PATH = Path("data/processed/scored_events.csv")
TIMELINE_FIGURE = Path("paper/figures/timeline_risk.png")
WINDOW_FIGURE = Path("paper/figures/intervention_window.png")
MANUSCRIPT_PATH = Path("paper/manuscript_assembled.md")


SUMMARY_ROWS = [
    ["Scenario", "Detected before export", "Intervention window", "Early signals"],
    ["Manual abuse", "Yes", "28.0 minutes", "abnormal login, unknown new device"],
    ["AI-assisted abuse", "Yes", "11.0 minutes", "abnormal login, reset-linked new device"],
]

DETECTOR_ROWS = [
    ["Detector", "Signal", "Purpose"],
    ["AbnormalLoginDetector", "login context", "early account-access anomaly"],
    ["HelpdeskResetDetector", "account recovery", "reset-workflow risk"],
    ["NewDeviceRiskDetector", "device enrollment", "account takeover or persistence signal"],
    ["OAuthScopeRiskDetector", "OAuth scopes", "risky application consent"],
    ["SaaSExportAnomalyDetector", "export volume", "late collection/export signal"],
]


def read_event_rows(limit: int = 12) -> list[list[str]]:
    if not SCORED_EVENTS_PATH.exists():
        return [["No scored event CSV found", "", "", "", ""]]

    rows: list[list[str]] = [["Scenario", "Event", "Event score", "Cumulative", "Detected"]]
    with SCORED_EVENTS_PATH.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        for index, row in enumerate(reader):
            if index >= limit:
                break
            rows.append(
                [
                    row.get("scenario", ""),
                    row.get("event_type", ""),
                    row.get("event_score", ""),
                    row.get("cumulative_score", ""),
                    row.get("detected", ""),
                ]
            )
    return rows


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        header_cells[index].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Identity Abuse Lab")
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Polished Research Artifact Report").italic = True

    document.add_paragraph(
        "A defensive synthetic lab for evaluating whether correlated identity and SaaS "
        "telemetry can detect simulated manual and AI-assisted account-abuse chains "
        "before data export."
    )

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "The current synthetic result set suggests that AI-assisted identity abuse may not "
        "require fundamentally different telemetry categories to detect, but it can compress "
        "the defender response window enough that slow manual review and late-stage export "
        "alerts become less reliable."
    )

    document.add_heading("Result Snapshot", level=1)
    add_docx_table(document, SUMMARY_ROWS)

    document.add_heading("Figures", level=1)
    if TIMELINE_FIGURE.exists():
        document.add_paragraph("Figure 1. Cumulative identity-abuse risk over scenario timeline.")
        document.add_picture(str(TIMELINE_FIGURE), width=Inches(6.2))
    if WINDOW_FIGURE.exists():
        document.add_paragraph("Figure 2. Intervention window before simulated SaaS export.")
        document.add_picture(str(WINDOW_FIGURE), width=Inches(6.2))

    document.add_heading("Detector Model", level=1)
    add_docx_table(document, DETECTOR_ROWS)

    document.add_heading("Event-Level Evidence Sample", level=1)
    add_docx_table(document, read_event_rows())

    document.add_heading("Interpretation", level=1)
    document.add_paragraph(
        "The enriched detector set moves detection earlier by using abnormal-login and "
        "new-device signals before OAuth consent or SaaS export. The AI-assisted scenario "
        "still leaves a shorter intervention window, making operational tempo a central "
        "defensive variable."
    )

    document.add_heading("Safety Boundary", level=1)
    document.add_paragraph(
        "This artifact is synthetic-only and defensive. It excludes real phishing, credential "
        "collection, malware, exploit development, unauthorized access guidance, production "
        "logs, real user data, customer data, and real incident-sensitive material."
    )

    document.save(DOCX_PATH)


def table_style() -> TableStyle:
    return TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
    )


def add_pdf_table(story: list, rows: list[list[str]]) -> None:
    table = Table(rows, repeatRows=1)
    table.setStyle(table_style())
    story.append(table)
    story.append(Spacer(1, 12))


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter)
    story: list = []

    story.append(Paragraph("Identity Abuse Lab", styles["Title"]))
    story.append(Paragraph("Polished Research Artifact Report", styles["Heading2"]))
    story.append(
        Paragraph(
            "A defensive synthetic lab for evaluating whether correlated identity and SaaS "
            "telemetry can detect simulated manual and AI-assisted account-abuse chains "
            "before data export.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    story.append(
        Paragraph(
            "The current synthetic result set suggests that AI-assisted identity abuse may not "
            "require fundamentally different telemetry categories to detect, but it can compress "
            "the defender response window enough that slow manual review and late-stage export "
            "alerts become less reliable.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Result Snapshot", styles["Heading1"]))
    add_pdf_table(story, SUMMARY_ROWS)

    if TIMELINE_FIGURE.exists() or WINDOW_FIGURE.exists():
        story.append(PageBreak())
        story.append(Paragraph("Figures", styles["Heading1"]))
        if TIMELINE_FIGURE.exists():
            story.append(Paragraph("Figure 1. Cumulative identity-abuse risk.", styles["BodyText"]))
            story.append(Image(str(TIMELINE_FIGURE), width=430, height=300))
            story.append(Spacer(1, 12))
        if WINDOW_FIGURE.exists():
            story.append(Paragraph("Figure 2. Intervention window before export.", styles["BodyText"]))
            story.append(Image(str(WINDOW_FIGURE), width=430, height=300))

    story.append(PageBreak())
    story.append(Paragraph("Detector Model", styles["Heading1"]))
    add_pdf_table(story, DETECTOR_ROWS)
    story.append(Paragraph("Event-Level Evidence Sample", styles["Heading1"]))
    add_pdf_table(story, read_event_rows())

    story.append(Paragraph("Safety Boundary", styles["Heading1"]))
    story.append(
        Paragraph(
            "This artifact is synthetic-only and defensive. It excludes real phishing, credential "
            "collection, malware, exploit development, unauthorized access guidance, production "
            "logs, real user data, customer data, and real incident-sensitive material.",
            styles["BodyText"],
        )
    )

    doc.build(story)


def build_zip() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in [DOCX_PATH, PDF_PATH, SCORED_EVENTS_PATH, TIMELINE_FIGURE, WINDOW_FIGURE, MANUSCRIPT_PATH]:
            if path.exists():
                archive.write(path, arcname=path.name)
        readme_path = OUTPUT_DIR / "README.txt"
        readme_path.write_text(
            "Identity Abuse Lab polished deliverable package.\n\n"
            "Includes polished DOCX/PDF reports, scored event CSV, generated figures, "
            "and assembled manuscript when available.\n",
            encoding="utf-8",
        )
        archive.write(readme_path, arcname="README.txt")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    build_zip()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {ZIP_PATH}")


if __name__ == "__main__":
    main()
