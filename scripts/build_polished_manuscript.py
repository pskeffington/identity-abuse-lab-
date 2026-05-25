from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUTPUT_DIR = Path("paper/polished")
ASSEMBLED_MANUSCRIPT = Path("paper/manuscript_assembled.md")
MANUSCRIPT_DOCX = OUTPUT_DIR / "Identity_Abuse_Lab_Manuscript.docx"
MANUSCRIPT_PDF = OUTPUT_DIR / "Identity_Abuse_Lab_Manuscript.pdf"


def read_manuscript() -> str:
    if not ASSEMBLED_MANUSCRIPT.exists():
        raise FileNotFoundError(
            "paper/manuscript_assembled.md not found. Run scripts/assemble_manuscript.py first."
        )
    return ASSEMBLED_MANUSCRIPT.read_text(encoding="utf-8")


def clean_heading(line: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped.startswith("#"):
        return None
    level = len(stripped) - len(stripped.lstrip("#"))
    text = stripped[level:].strip()
    if not text:
        return None
    return min(level, 3), text


def build_docx(markdown_text: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Identity Abuse Lab")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Manuscript Draft").italic = True

    document.add_paragraph(
        "Synthetic defensive research on AI-assisted identity-abuse detection, timeline "
        "correlation, and response-window compression."
    )

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        heading = clean_heading(line)
        if heading:
            level, text = heading
            document.add_heading(text, level=level)
            continue
        if line.startswith("|"):
            # Keep markdown tables readable without trying to reconstruct every table shape.
            para = document.add_paragraph(line)
            para.style = document.styles["No Spacing"]
            continue
        para = document.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)

    document.save(MANUSCRIPT_DOCX)


def paragraph_style_for_heading(styles, level: int):
    if level == 1:
        return styles["Heading1"]
    if level == 2:
        return styles["Heading2"]
    return styles["Heading3"]


def build_pdf(markdown_text: str) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(MANUSCRIPT_PDF), pagesize=letter)
    story: list = []

    story.append(Paragraph("Identity Abuse Lab", styles["Title"]))
    story.append(Paragraph("Manuscript Draft", styles["Heading2"]))
    story.append(
        Paragraph(
            "Synthetic defensive research on AI-assisted identity-abuse detection, timeline "
            "correlation, and response-window compression.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = clean_heading(line)
        if heading:
            level, text = heading
            story.append(Spacer(1, 8))
            story.append(Paragraph(text, paragraph_style_for_heading(styles, level)))
            continue
        if line.startswith("|"):
            story.append(Paragraph(line.replace("|", " | "), styles["Code"]))
            continue
        story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 4))

    doc.build(story)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manuscript = read_manuscript()
    build_docx(manuscript)
    build_pdf(manuscript)
    print(f"Wrote {MANUSCRIPT_DOCX}")
    print(f"Wrote {MANUSCRIPT_PDF}")


if __name__ == "__main__":
    main()
